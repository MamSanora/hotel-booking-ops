import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_document():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Hotel Booking Ops: Development Update Summary', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Introduction
    doc.add_paragraph(
        'This document provides a comprehensive summary of the development work and refactoring '
        'carried out during the recent engineering session. The primary focus areas included authentication enhancements, '
        'database normalization, extensive codebase refactoring, and UI improvements for the payment gateway.'
    )
    
    # 1. Guest Password Recovery
    doc.add_heading('1. Guest Password Recovery Implementation', level=1)
    p = doc.add_paragraph()
    p.add_run('Overview: ').bold = True
    p.add_run('Implemented a secure guest password recovery workflow.')
    
    doc.add_paragraph('Key Points:', style='List Bullet')
    doc.add_paragraph('Established a primary recovery mechanism for guest accounts.', style='List Bullet 2')
    doc.add_paragraph('Analyzed the viability of a fallback OTP via phone number. It was determined that multiple test accounts share the same phone number in the existing database. To avoid security risks and potential data anomalies, a fallback OTP system was excluded by design.', style='List Bullet 2')
    
    # 2. Database Normalization (Rooms)
    doc.add_heading('2. Database Normalization (Rooms)', level=1)
    p = doc.add_paragraph()
    p.add_run('Overview: ').bold = True
    p.add_run('Resolved a significant data anomaly (update anomaly) in the rooms table by extracting duplicated data into a centralized table.')
    
    doc.add_paragraph('Key Changes:', style='List Bullet')
    doc.add_paragraph('Created a new table: ', style='List Bullet 2').add_run('room_types').bold = True
    doc.add_paragraph('Migrated room characteristics (price_per_night, capacity, description) to the new table, preventing inconsistent pricing data across identical room types.', style='List Bullet 2')
    doc.add_paragraph('Added a ', style='List Bullet 2').add_run('room_type_id').bold = True
    doc.add_paragraph(' foreign key to the rooms table.', style='List Bullet 2')
    doc.add_paragraph('Dropped redundant columns from the rooms table to achieve a normalized relational structure.', style='List Bullet 2')
    doc.add_paragraph('Updated the RoomSeeder to correctly populate the normalized schema.', style='List Bullet 2')

    # 3. Code Refactoring
    doc.add_heading('3. Codebase Refactoring (Controllers & Views)', level=1)
    p = doc.add_paragraph()
    p.add_run('Overview: ').bold = True
    p.add_run('Updated the application layer to reflect the new database structure, ensuring that room pricing and details are correctly retrieved via Eloquent relationships.')
    
    doc.add_paragraph('Models:', style='List Bullet')
    doc.add_paragraph('Updated the Room model to include a belongsTo(RoomType::class) relationship.', style='List Bullet 2')
    
    doc.add_paragraph('Controllers:', style='List Bullet')
    doc.add_paragraph('AdminRoomController: Rewrote form handling to use room_type_id dropdowns. Replaced StoreRoomRequest with inline validation logic.', style='List Bullet 2')
    doc.add_paragraph('Guest\\RoomController, ReceptionDashboardController, WalkInBookingController: Updated all pricing calculators to reference $room->roomType->price_per_night.', style='List Bullet 2')
    
    doc.add_paragraph('Views:', style='List Bullet')
    doc.add_paragraph('Updated 9+ blade templates across the application (guest homes, room listings, invoices, admin panels, reception dashboards) to dynamically pull data from the new relationship.', style='List Bullet 2')

    # 4. UI Enhancements
    doc.add_heading('4. UI Enhancements (KHQR Payment Page)', level=1)
    p = doc.add_paragraph()
    p.add_run('Overview: ').bold = True
    p.add_run('Fixed visual bugs and aligned the payment gateway interface with the target design mockup.')
    
    doc.add_paragraph('Key Fixes:', style='List Bullet')
    doc.add_paragraph('Resolved a severe CSS layout conflict where a "flex items-center" wrapper forced the grid columns into a narrow, squished layout.', style='List Bullet 2')
    doc.add_paragraph('Restructured the CSS Grid to a 5/12 left (KHQR Card) and 7/12 right (Info Panels) split.', style='List Bullet 2')
    doc.add_paragraph('Redesigned the "Supported Banking Apps" section to display as a clean horizontal strip of 4 icons.', style='List Bullet 2')
    doc.add_paragraph('Refined typography, background colors, card shadows, and instructional step numbers to create a premium, polished user experience.', style='List Bullet 2')

    # Save
    filename = 'Development_Update_Summary.docx'
    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == '__main__':
    create_document()
